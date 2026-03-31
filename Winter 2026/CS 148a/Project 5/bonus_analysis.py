#!/usr/bin/env python3
from __future__ import annotations

import gc
import io
import json
import math
import platform
import subprocess
import textwrap
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
from PIL import Image
from docx import Document
from docx.shared import Inches
from huggingface_hub import hf_hub_download
from torchvision import transforms


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "bonus_outputs"
FIG_DIR = OUT_DIR / "figures"
OUT_DIR.mkdir(exist_ok=True)
FIG_DIR.mkdir(exist_ok=True)

TO_TENSOR = transforms.ToTensor()
DEVICE = torch.device("cpu")
SEED = 42
TRAIN_FRACTION = 0.85
VAL_TIMING_IMAGES = 16
PIPELINE_BATCH_SIZE = 32
PIPELINE_PRESTACK_SIZE = 256

PROJECT4_NOTEBOOK = ROOT.parent / "Project 4" / "CS148a_proj4_FM_starter.ipynb"
PROJECT3_TUNING_JSON = ROOT.parent / "Project 3" / "tuning_results.json"


PROJECT4_METRICS = {
    "CLIP zero-shot": {
        "family": "Foundation",
        "mode": "Zero-shot",
        "train_samples": math.nan,
        "train_acc": math.nan,
        "val_acc": 0.548220,
        "params": 427_944_193,
        "seconds_per_image": math.nan,
        "notes": "Project 4 notebook output; full 1545-image validation split.",
    },
    "CLIP downstream": {
        "family": "Foundation",
        "mode": "Downstream MLP",
        "train_samples": 17_498,
        "train_acc": math.nan,
        "val_acc": 0.810356,
        "params": 427_944_193,
        "seconds_per_image": math.nan,
        "notes": "Project 4 notebook output; CLIP encoder plus learned MLP head.",
    },
    "DINO downstream": {
        "family": "Foundation",
        "mode": "Downstream MLP",
        "train_samples": 17_498,
        "train_acc": math.nan,
        "val_acc": 0.680906,
        "params": 304_368_640,
        "seconds_per_image": math.nan,
        "notes": "Project 4 notebook output; DINOv2 encoder plus learned MLP head.",
    },
    "CLIP + DINO fusion": {
        "family": "Foundation",
        "mode": "Feature fusion",
        "train_samples": 17_498,
        "train_acc": math.nan,
        "val_acc": 0.828479,
        "params": 427_944_193 + 304_368_640,
        "seconds_per_image": math.nan,
        "notes": "Project 4 notebook output; concatenated CLIP and DINO features with fusion head.",
    },
    "Qwen2.5-VL zero-shot": {
        "family": "Foundation",
        "mode": "Autoregressive VLM",
        "train_samples": math.nan,
        "train_acc": math.nan,
        "val_acc": 0.316406,
        "params": 3_754_622_976,
        "seconds_per_image": math.nan,
        "notes": "Project 4 notebook output; accuracy reported on 256 validation images, not the full 1545.",
    },
}


@dataclass
class EvalResult:
    train_acc: float
    val_acc: float
    params: int
    seconds_per_image: float


def count_parameters(module: torch.nn.Module) -> int:
    return sum(p.numel() for p in module.parameters())


def compute_mlp_params(input_dim: int, hidden_dims: tuple[int, int], output_dim: int = 10) -> int:
    h1, h2 = hidden_dims
    params = 2 * input_dim  # LayerNorm weight + bias
    params += input_dim * h1 + h1
    params += h1 * h2 + h2
    params += h2 * output_dim + output_dim
    return params


def build_feature_mlp(input_dim: int, hidden_dims: tuple[int, int], dropout: float) -> torch.nn.Module:
    h1, h2 = hidden_dims
    return torch.nn.Sequential(
        torch.nn.LayerNorm(input_dim),
        torch.nn.Linear(input_dim, h1),
        torch.nn.GELU(),
        torch.nn.Dropout(dropout),
        torch.nn.Linear(h1, h2),
        torch.nn.GELU(),
        torch.nn.Dropout(dropout),
        torch.nn.Linear(h2, 10),
    )


def safe_cpu_name() -> str:
    try:
        value = subprocess.check_output(
            ["sysctl", "-n", "machdep.cpu.brand_string"],
            text=True,
        ).strip()
        if value:
            return value
    except Exception:
        pass
    return f"{platform.system()} {platform.machine()}"


def decode_image(record: dict) -> Image.Image:
    image_field = record["image"]
    raw_bytes = image_field.get("bytes")
    if raw_bytes is None:
        raise ValueError("Dataset row is missing image bytes.")
    with Image.open(io.BytesIO(raw_bytes)) as img:
        return img.convert("RGB")


def load_dataset_records() -> list[dict]:
    parquet_path = hf_hub_download(
        repo_id="EE148-project/MNIST-in-the-world",
        filename="data/train-00000-of-00001.parquet",
        repo_type="dataset",
    )
    df = pd.read_parquet(parquet_path)
    return df.to_dict(orient="records")


def make_split_indices(num_samples: int) -> tuple[list[int], list[int]]:
    num_train = int(TRAIN_FRACTION * num_samples)
    perm = torch.randperm(num_samples, generator=torch.Generator().manual_seed(SEED)).tolist()
    train_indices = perm[:num_train]
    val_indices = perm[num_train:]
    return train_indices, val_indices


def load_scripted_pipeline(filename: str) -> torch.jit.ScriptModule:
    path = hf_hub_download(
        repo_id="osanan/ee148a-project",
        filename=filename,
        repo_type="model",
    )
    model = torch.jit.load(path, map_location="cpu")
    model.device = torch.device("cpu")
    model.eval()
    return model


def batch_tensors(records: list[dict], indices: Iterable[int]) -> torch.Tensor:
    tensors = [
        TO_TENSOR(
            decode_image(records[idx]).resize(
                (PIPELINE_PRESTACK_SIZE, PIPELINE_PRESTACK_SIZE),
                Image.Resampling.BILINEAR,
            )
        )
        for idx in indices
    ]
    return torch.stack(tensors)


def evaluate_scripted_pipeline(
    model: torch.jit.ScriptModule,
    records: list[dict],
    labels: np.ndarray,
    indices: list[int],
    batch_size: int = PIPELINE_BATCH_SIZE,
) -> float:
    preds: list[int] = []
    with torch.no_grad():
        for start in range(0, len(indices), batch_size):
            batch_indices = indices[start : start + batch_size]
            xb = batch_tensors(records, batch_indices)
            out = model(xb)
            preds.extend(int(x) for x in out.cpu().tolist())
    target = labels[np.array(indices)]
    return float((np.array(preds) == target).mean())


def time_scripted_pipeline(
    model: torch.jit.ScriptModule,
    records: list[dict],
    indices: list[int],
    batch_size: int = PIPELINE_BATCH_SIZE,
    repeats: int = 3,
) -> float:
    probe_indices = indices[: max(1, min(len(indices), VAL_TIMING_IMAGES))]
    batches: list[torch.Tensor] = []
    for start in range(0, len(probe_indices), batch_size):
        batch_idx = probe_indices[start : start + batch_size]
        batches.append(batch_tensors(records, batch_idx))

    if not batches:
        return math.nan

    with torch.no_grad():
        for xb in batches:
            _ = model(xb)

    timings = []
    num_images = sum(batch.shape[0] for batch in batches)
    with torch.no_grad():
        for _ in range(repeats):
            t0 = time.perf_counter()
            for xb in batches:
                _ = model(xb)
            timings.append((time.perf_counter() - t0) / num_images)
    return float(np.median(timings))


def evaluate_exported_models(records: list[dict], labels: np.ndarray, train_indices: list[int], val_indices: list[int]) -> dict[str, EvalResult]:
    results: dict[str, EvalResult] = {}
    model_specs = {
        "CNN": "pipeline-cnn.pt",
        "ViT": "pipeline-vit.pt",
    }

    for name, filename in model_specs.items():
        model = load_scripted_pipeline(filename)
        params = count_parameters(model)
        sec_per_image = time_scripted_pipeline(model, records, val_indices)
        results[name] = EvalResult(
            train_acc=math.nan,
            val_acc=math.nan,
            params=params,
            seconds_per_image=sec_per_image,
        )
        del model
        gc.collect()

    return results


def measure_clip_timings(records: list[dict], indices: list[int]) -> dict[str, float]:
    from transformers import CLIPModel, CLIPProcessor

    def unwrap_features(obj):
        if torch.is_tensor(obj):
            return obj
        for attr in ("image_embeds", "text_embeds", "pooler_output"):
            value = getattr(obj, attr, None)
            if value is not None:
                return value
        raise TypeError(f"Unsupported CLIP feature object: {type(obj).__name__}")

    model_id = "openai/clip-vit-large-patch14-336"
    processor = CLIPProcessor.from_pretrained(model_id)
    model = CLIPModel.from_pretrained(model_id)
    model.eval()

    images = [decode_image(records[idx]) for idx in indices[:VAL_TIMING_IMAGES]]
    prompts = [f"a photo of the digit {i}" for i in range(10)]
    text_inputs = processor(text=prompts, return_tensors="pt", padding=True)
    with torch.no_grad():
        text_features = unwrap_features(model.get_text_features(**text_inputs))
        text_features = text_features / text_features.norm(dim=-1, keepdim=True)

    head = build_feature_mlp(768, (1024, 512), 0.20).eval()

    def zero_shot_step() -> None:
        inputs = processor(images=images, return_tensors="pt")
        with torch.no_grad():
            image_features = unwrap_features(model.get_image_features(pixel_values=inputs["pixel_values"]))
            image_features = image_features / image_features.norm(dim=-1, keepdim=True)
            _ = image_features @ text_features.T

    def downstream_step() -> None:
        inputs = processor(images=images, return_tensors="pt")
        with torch.no_grad():
            image_features = unwrap_features(model.get_image_features(pixel_values=inputs["pixel_values"]))
            _ = head(image_features)

    def measure(step) -> float:
        step()
        runs = []
        for _ in range(3):
            t0 = time.perf_counter()
            step()
            runs.append((time.perf_counter() - t0) / len(images))
        return float(np.median(runs))

    zero_shot = measure(zero_shot_step)
    downstream = measure(downstream_step)

    del head
    del model
    gc.collect()
    return {
        "CLIP zero-shot": zero_shot,
        "CLIP downstream": downstream,
    }


def measure_dino_timing(records: list[dict], indices: list[int]) -> float:
    from transformers import AutoImageProcessor, AutoModel

    model_id = "facebook/dinov2-large"
    processor = AutoImageProcessor.from_pretrained(model_id)
    model = AutoModel.from_pretrained(model_id)
    model.eval()
    head = build_feature_mlp(1024, (1536, 768), 0.15).eval()

    images = [decode_image(records[idx]) for idx in indices[:VAL_TIMING_IMAGES]]

    def step() -> None:
        inputs = processor(images=images, return_tensors="pt")
        with torch.no_grad():
            features = model(pixel_values=inputs["pixel_values"]).last_hidden_state[:, 0, :]
            _ = head(features)

    step()
    runs = []
    for _ in range(3):
        t0 = time.perf_counter()
        step()
        runs.append((time.perf_counter() - t0) / len(images))

    del head
    del model
    gc.collect()
    return float(np.median(runs))


def build_results_dataframe(exported: dict[str, EvalResult], timing_notes: list[str]) -> pd.DataFrame:
    rows = [
        {
            "model": "CNN",
            "family": "CNN",
            "mode": "Submitted TorchScript pipeline",
            "train_samples": 8_749,
            "train_acc": math.nan,
            "val_acc": 0.871197,
            "params": exported["CNN"].params,
            "seconds_per_image": exported["CNN"].seconds_per_image,
            "notes": "Accuracy sourced from Project 2 writeup PDF (best validation accuracy 87.1197%); params/timing from osanan/ee148a-project pipeline-cnn.pt.",
        },
        {
            "model": "ViT",
            "family": "ViT",
            "mode": "Submitted TorchScript pipeline",
            "train_samples": 8_749,
            "train_acc": 0.9693,
            "val_acc": 0.9625,
            "params": exported["ViT"].params,
            "seconds_per_image": exported["ViT"].seconds_per_image,
            "notes": "Accuracy sourced from Project 3 writeup PDF (train no-aug 96.93%, val 96.25%); params/timing from osanan/ee148a-project pipeline-vit.pt. The sample-count axis still excludes the additional 60K-image MNIST pretraining stage.",
        },
    ]

    clip_head_params = compute_mlp_params(768, (1024, 512))
    dino_head_params = compute_mlp_params(1024, (1536, 768))
    fusion_head_params = compute_mlp_params(768 + 1024, (1536, 768))

    params_overrides = {
        "CLIP downstream": PROJECT4_METRICS["CLIP downstream"]["params"] + clip_head_params,
        "DINO downstream": PROJECT4_METRICS["DINO downstream"]["params"] + dino_head_params,
        "CLIP + DINO fusion": PROJECT4_METRICS["CLIP + DINO fusion"]["params"] + fusion_head_params,
    }

    for model_name, info in PROJECT4_METRICS.items():
        rows.append(
            {
                "model": model_name,
                "family": info["family"],
                "mode": info["mode"],
                "train_samples": info["train_samples"],
                "train_acc": info["train_acc"],
                "val_acc": info["val_acc"],
                "params": params_overrides.get(model_name, info["params"]),
                "seconds_per_image": info["seconds_per_image"],
                "notes": info["notes"],
            }
        )

    df = pd.DataFrame(rows)
    df["train_error"] = 1.0 - df["train_acc"]
    df["val_error"] = 1.0 - df["val_acc"]
    if timing_notes:
        notes_path = OUT_DIR / "timing_notes.txt"
        notes_path.write_text("\n".join(timing_notes) + "\n")
    return df


def annotate_points(ax, subset: pd.DataFrame, x: str, y: str, fontsize: int = 8) -> None:
    for _, row in subset.iterrows():
        ax.annotate(
            row["model"],
            (row[x], row[y]),
            textcoords="offset points",
            xytext=(4, 4),
            fontsize=fontsize,
        )


def plot_log_error_vs_sample_count(df: pd.DataFrame) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.8))

    val_subset = df[df["train_samples"].notna() & df["val_error"].notna()].copy()
    train_subset = df[df["train_samples"].notna() & df["train_error"].notna()].copy()

    axes[0].scatter(val_subset["train_samples"], val_subset["val_error"], s=65, color="#15616d")
    annotate_points(axes[0], val_subset, "train_samples", "val_error")
    axes[0].set_xscale("log")
    axes[0].set_yscale("log")
    axes[0].set_xlabel("Effective downstream sample count")
    axes[0].set_ylabel("Validation error")
    axes[0].set_title("Validation error vs. sample count")
    axes[0].grid(True, which="both", alpha=0.25)

    axes[1].scatter(train_subset["train_samples"], train_subset["train_error"], s=65, color="#8d0801")
    annotate_points(axes[1], train_subset, "train_samples", "train_error")
    axes[1].set_xscale("log")
    axes[1].set_yscale("log")
    axes[1].set_xlabel("Effective downstream sample count")
    axes[1].set_ylabel("Training error")
    axes[1].set_title("Training error vs. sample count")
    axes[1].grid(True, which="both", alpha=0.25)

    fig.suptitle("Plot 1: log(error) vs. log(sample count)", fontsize=14)
    fig.tight_layout(rect=(0, 0, 1, 0.95))
    out_path = FIG_DIR / "plot1_log_error_vs_sample_count.png"
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_log_error_vs_parameter_count(df: pd.DataFrame) -> Path:
    subset = df[df["params"].notna() & df["val_error"].notna()].copy()
    fig, ax = plt.subplots(figsize=(7.6, 5.2))
    ax.scatter(subset["params"], subset["val_error"], s=72, color="#0a9396")
    annotate_points(ax, subset, "params", "val_error")
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("Parameter count")
    ax.set_ylabel("Validation error")
    ax.set_title("Plot 2: log(error) vs. log(parameter count)")
    ax.grid(True, which="both", alpha=0.25)
    fig.tight_layout()
    out_path = FIG_DIR / "plot2_log_error_vs_parameter_count.png"
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_log_error_vs_wall_clock(df: pd.DataFrame) -> Path:
    subset = df[df["seconds_per_image"].notna() & df["val_error"].notna()].copy()
    fig, ax = plt.subplots(figsize=(7.6, 5.2))
    ax.scatter(subset["seconds_per_image"], subset["val_error"], s=72, color="#bb3e03")
    annotate_points(ax, subset, "seconds_per_image", "val_error")
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("Seconds per image")
    ax.set_ylabel("Validation error")
    ax.set_title("Plot 3: log(error) vs. log(wall-clock inference time)")
    ax.grid(True, which="both", alpha=0.25)
    fig.tight_layout()
    out_path = FIG_DIR / "plot3_log_error_vs_wall_clock.png"
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out_path


def plot_vit_sweep() -> Path:
    with PROJECT3_TUNING_JSON.open() as f:
        tuning = json.load(f)

    df = pd.DataFrame(tuning)
    fig, ax = plt.subplots(figsize=(8, 5.2))
    ax.scatter(df["gap"], df["val_acc"], s=68, color="#5f0f40")
    for _, row in df.iterrows():
        ax.annotate(row["name"], (row["gap"], row["val_acc"]), textcoords="offset points", xytext=(4, 4), fontsize=8)
    ax.set_xlabel("Train-validation gap (%)")
    ax.set_ylabel("Validation accuracy (%)")
    ax.set_title("Plot N: Project 3 ViT tuning sweep")
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    out_path = FIG_DIR / "plotN_vit_tuning_gap_vs_val_accuracy.png"
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out_path


def add_table(document: Document, df: pd.DataFrame) -> None:
    cols = ["Model", "Params", "Train samples", "Train acc", "Val acc", "Sec/image", "Notes"]
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for idx, name in enumerate(cols):
        table.rows[0].cells[idx].text = name

    for _, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["model"])
        cells[1].text = f"{int(row['params']):,}" if pd.notna(row["params"]) else "n/a"
        cells[2].text = f"{int(row['train_samples']):,}" if pd.notna(row["train_samples"]) else "n/a"
        cells[3].text = f"{100 * row['train_acc']:.2f}%" if pd.notna(row["train_acc"]) else "n/a"
        cells[4].text = f"{100 * row['val_acc']:.2f}%" if pd.notna(row["val_acc"]) else "n/a"
        cells[5].text = f"{row['seconds_per_image']:.4f}" if pd.notna(row["seconds_per_image"]) else "n/a"
        cells[6].text = textwrap.shorten(str(row["notes"]), width=95, placeholder="...")


def build_docx(df: pd.DataFrame, figure_paths: list[Path], timing_notes: list[str], compute_resource: str) -> Path:
    doc = Document()
    doc.add_heading("CS 148a Project 5 Writeup", level=0)

    doc.add_paragraph(
        "This writeup compares the finalized Project 2, Project 3, and Project 4 checkpoints using the bonus-assignment plots: "
        "error vs. sample count, parameter count, and inference wall clock, plus one additional plot based on the ViT tuning sweep."
    )

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Canonical artifacts used:\n"
        "- Project 2 reported accuracy source: `../Project 2/CS 148a Project 2 Writeup.pdf`\n"
        "- Project 3 reported accuracy source: `../Project 3/CS 148a Project 3 Writeup.pdf`\n"
        "- Project 2 submitted checkpoint for params/timing: osanan/ee148a-project `pipeline-cnn.pt`\n"
        "- Project 3 submitted checkpoint for params/timing: osanan/ee148a-project `pipeline-vit.pt`\n"
        "- Project 4 metrics: extracted from the executed notebook `../Project 4/CS148a_proj4_FM_starter.ipynb`\n"
        "- Additional Project 3 tuning data: `../Project 3/tuning_results.json`"
    )
    doc.add_paragraph(
        f"Dataset and split: `MNIST-in-the-world` from Hugging Face, 10,294 total images, 8,749 train and 1,545 validation examples using a fixed seed-42 85/15 split."
    )
    doc.add_paragraph(
        f"Compute resource for wall-clock measurements: {compute_resource}. This environment did not expose CUDA, so all wall-clock values in this writeup are CPU-side measurements."
    )
    doc.add_paragraph(
        "Accuracy sourcing policy: Project 2 and Project 3 accuracy values are taken from the submitted writeup PDFs, because those are the reported course-submission numbers. "
        "Their exported Hugging Face pipelines are used only for parameter counting and local timing."
    )
    doc.add_paragraph(
        "Caveats: zero-shot methods have no meaningful downstream sample count, so they are excluded from the sample-count plot; "
        "Qwen accuracy in Project 4 was only reported on 256 validation images; "
        "the Project 3 notebook describes an additional 60K-image MNIST pretraining phase, but the sample-count axis here only tracks downstream labeled samples on the target dataset."
    )
    if timing_notes:
        doc.add_paragraph("Timing notes:\n" + "\n".join(f"- {note}" for note in timing_notes))

    doc.add_heading("Compiled metrics", level=1)
    add_table(doc, df)

    best_val = df.loc[df["val_acc"].idxmax()]
    fastest = df.loc[df["seconds_per_image"].fillna(np.inf).idxmin()]
    smallest = df.loc[df["params"].idxmin()]

    doc.add_heading("Summary observations", level=1)
    doc.add_paragraph(
        f"The strongest validation result in this comparison is {best_val['model']} at {100 * best_val['val_acc']:.2f}% validation accuracy."
    )
    doc.add_paragraph(
        f"The smallest exported submitted model is {smallest['model']} with {int(smallest['params']):,} parameters."
    )
    if pd.notna(fastest["seconds_per_image"]):
        doc.add_paragraph(
            f"The fastest measured inference path is {fastest['model']} at {fastest['seconds_per_image']:.4f} seconds per image on CPU."
        )
    doc.add_paragraph(
        "Parameter scaling is not monotonic: much larger foundation models do not automatically win on this dataset unless their representations are paired with a strong downstream classifier. "
        "The biggest gain in Project 4 comes from moving from CLIP zero-shot to CLIP plus a learned head, which cuts validation error substantially without changing the frozen backbone."
    )
    doc.add_paragraph(
        "The submitted ViT remains the strongest fully supervised exported checkpoint in the set, but that result depends on more training sophistication than the simple sample-count axis captures: stronger augmentation, a hybrid tokenizer, and an extra clean-MNIST pretraining stage."
    )

    captions = [
        "Figure 1. Required Plot 1. Validation and training error against effective downstream sample count. Zero-shot models are excluded because log(0) is undefined.",
        "Figure 2. Required Plot 2. Validation error against parameter count across the submitted CNN, submitted ViT, and Project 4 foundation-model variants.",
        "Figure 3. Required Plot 3. Validation error against measured CPU inference time per image. Qwen is omitted from this plot because no comparable CPU timing run was performed in this environment.",
        "Figure 4. Additional plot. Project 3 ViT tuning sweep showing that a smaller train-validation gap did not always produce the best validation accuracy.",
    ]

    doc.add_heading("Figures", level=1)
    for path, caption in zip(figure_paths, captions):
        doc.add_picture(str(path), width=Inches(6.6))
        doc.add_paragraph(caption)

    doc.add_heading("Code / artifact links", level=1)
    doc.add_paragraph(
        "Relevant local files for upload/linking:\n"
        "- ./bonus_analysis.py\n"
        "- ./bonus_outputs/compiled_metrics.csv\n"
        "- ../Project 2/CS148a_proj2_CNN_starter.ipynb and ../Project 2/CS148a_proj2_CNN_starter_resnet.ipynb\n"
        "- ../Project 3/CS148a_proj3_ViT_starter_completed.ipynb\n"
        "- ../Project 4/CS148a_proj4_FM_starter.ipynb\n"
        "- Hugging Face submitted model repo: https://huggingface.co/osanan/ee148a-project"
    )

    doc.add_heading("AI usage", level=1)
    doc.add_paragraph(
        "This analysis and draft writeup were prepared with OpenAI Codex assistance in the local workspace. "
        "Per the assignment instructions, the relevant conversation log should be attached or copied into the final submission materials."
    )

    out_path = OUT_DIR / "CS 148a Project 5 Writeup.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    torch.set_num_threads(max(1, min(8, torch.get_num_threads())))
    records = load_dataset_records()
    labels = np.array([int(r["label"]) for r in records], dtype=np.int64)
    train_indices, val_indices = make_split_indices(len(records))

    exported_results = evaluate_exported_models(records, labels, train_indices, val_indices)
    timing_notes: list[str] = []

    try:
        clip_timings = measure_clip_timings(records, val_indices)
        PROJECT4_METRICS["CLIP zero-shot"]["seconds_per_image"] = clip_timings["CLIP zero-shot"]
        PROJECT4_METRICS["CLIP downstream"]["seconds_per_image"] = clip_timings["CLIP downstream"]
    except Exception as exc:
        timing_notes.append(f"CLIP timing unavailable: {type(exc).__name__}: {exc}")

    try:
        PROJECT4_METRICS["DINO downstream"]["seconds_per_image"] = measure_dino_timing(records, val_indices)
    except Exception as exc:
        timing_notes.append(f"DINO timing unavailable: {type(exc).__name__}: {exc}")

    if pd.notna(PROJECT4_METRICS["CLIP downstream"]["seconds_per_image"]) and pd.notna(PROJECT4_METRICS["DINO downstream"]["seconds_per_image"]):
        PROJECT4_METRICS["CLIP + DINO fusion"]["seconds_per_image"] = (
            PROJECT4_METRICS["CLIP downstream"]["seconds_per_image"]
            + PROJECT4_METRICS["DINO downstream"]["seconds_per_image"]
        )
        timing_notes.append("Fusion wall-clock is approximated as CLIP downstream time plus DINO downstream time on CPU.")

    df = build_results_dataframe(exported_results, timing_notes)
    csv_path = OUT_DIR / "compiled_metrics.csv"
    df.to_csv(csv_path, index=False)

    fig1 = plot_log_error_vs_sample_count(df)
    fig2 = plot_log_error_vs_parameter_count(df)
    fig3 = plot_log_error_vs_wall_clock(df)
    fig4 = plot_vit_sweep()

    compute_resource = safe_cpu_name()
    docx_path = build_docx(df, [fig1, fig2, fig3, fig4], timing_notes, compute_resource)

    summary = {
        "dataset_size": len(records),
        "train_size": len(train_indices),
        "val_size": len(val_indices),
        "compute_resource": compute_resource,
        "outputs": {
            "metrics_csv": str(csv_path),
            "docx": str(docx_path),
            "figures": [str(fig1), str(fig2), str(fig3), str(fig4)],
        },
    }
    (OUT_DIR / "run_summary.json").write_text(json.dumps(summary, indent=2) + "\n")

    print(f"Wrote metrics to {csv_path}")
    print(f"Wrote writeup to {docx_path}")
    print("Figures:")
    for path in [fig1, fig2, fig3, fig4]:
        print(f"  - {path}")


if __name__ == "__main__":
    main()
